import base64
import csv
import io
import json
import os
import re
import tempfile
from datetime import datetime
from html import unescape
from pathlib import Path
from typing import Any
from xml.etree import ElementTree
from contextlib import contextmanager

from filelock import FileLock

from dotenv import load_dotenv
from greennode_agentbase import GreenNodeAgentBaseApp, PingStatus, RequestContext
from starlette.requests import Request
from starlette.responses import HTMLResponse, JSONResponse, Response

load_dotenv()

app = GreenNodeAgentBaseApp()

DATA_DIR = Path(".agentbase")
KNOWLEDGE_FILE = DATA_DIR / "knowledge.json"
MAX_CONTEXT_CHARS = 6000
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 180
TOP_CONTEXT_CHUNKS = 10
DEFAULT_MAX_UPLOAD_BYTES = 10 * 1024 * 1024

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".xml",
    ".svg",
    ".drawio",
    ".bpmn",
    ".mmd",
    ".mermaid",
    ".puml",
    ".plantuml",
    ".pdf",
    ".docx",
    ".xlsx",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
READY = "READY"
NEEDS_REVIEW = "NEEDS_REVIEW"
FAILED = "FAILED"
AUTO_LEARNED_TYPE = "auto-learned"

TESTCASE_TYPES_FILE = DATA_DIR / "testcase_types.json"

DEFAULT_TEST_CASE_TYPES: list[dict] = [
    {"type": "Positive", "priority": "High", "enabled": True, "condition": None},
    {"type": "Negative", "priority": "High", "enabled": True, "condition": None},
    {"type": "Boundary", "priority": "Medium", "enabled": True, "condition": None},
    {"type": "Permission", "priority": "High", "enabled": True, "condition": None},
    {"type": "Resilience", "priority": "Medium", "enabled": True, "condition": None},
    {"type": "Workflow", "priority": "High", "enabled": True, "condition": "has_workflow"},
]

QUALITY_CHECKLIST = [
    "Each case has a clear objective and expected result.",
    "Preconditions and test data are explicit.",
    "Positive, negative, boundary, workflow, integration, and permission scenarios are considered.",
    "Steps are executable by another QC without hidden assumptions.",
    "Priority and test type are assigned consistently.",
    "Generated cases reference the uploaded business documents where relevant.",
]

CHAT_HTML = """<!doctype html>
<html lang="vi">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>QC Test Case Agent</title>
  <style>
    :root {
      color-scheme: light;
      --bg: #eef1f8;
      --panel: #ffffff;
      --text: #1a2233;
      --muted: #6b7385;
      --line: #e3e8f0;
      --brand: #0d9488;
      --brand-dark: #0f766e;
      --brand-2: #6366f1;
      --soft: #eef7f6;
      --warn: #fff7e6;
      --danger: #b42318;
      --shadow: 0 1px 3px rgba(16,24,40,.06), 0 1px 2px rgba(16,24,40,.04);
      --shadow-lg: 0 14px 34px -14px rgba(16,24,40,.22);
      --radius: 14px;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      min-height: 100vh;
      background:
        radial-gradient(1100px 480px at -8% -10%, #d9f3ef 0%, transparent 55%),
        radial-gradient(1000px 500px at 112% -4%, #e7e8ff 0%, transparent 52%),
        var(--bg);
      color: var(--text);
      font: 15px/1.55 Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      -webkit-font-smoothing: antialiased;
    }
    ::-webkit-scrollbar { width: 10px; height: 10px; }
    ::-webkit-scrollbar-thumb { background: #cfd6e4; border-radius: 999px; border: 2px solid transparent; background-clip: padding-box; }
    ::-webkit-scrollbar-thumb:hover { background: #b6bfd1; background-clip: padding-box; }
    .shell {
      min-height: 100vh;
      display: grid;
      grid-template-columns: minmax(340px, 440px) 1fr;
    }
    aside {
      background: rgba(255,255,255,.82);
      backdrop-filter: blur(6px);
      border-right: 1px solid var(--line);
      padding: 18px;
      display: flex;
      flex-direction: column;
      gap: 14px;
      overflow: auto;
    }
    main { padding: 24px 28px; overflow: auto; }
    .hero {
      background: linear-gradient(135deg, var(--brand) 0%, var(--brand-2) 135%);
      color: #fff;
      border-radius: var(--radius);
      padding: 20px;
      box-shadow: var(--shadow-lg);
    }
    .hero h1 { margin: 0; font-size: 21px; font-weight: 800; display: flex; align-items: center; gap: 10px; }
    .hero .subtitle { margin: 8px 0 0; color: rgba(255,255,255,.88); font-size: 13px; }
    h2 {
      margin: 0 0 12px;
      font-size: 15px;
      font-weight: 800;
      display: flex;
      align-items: center;
      gap: 8px;
    }
    .subtitle { margin: 6px 0 0; color: var(--muted); }
    .panel {
      border: 1px solid var(--line);
      border-radius: var(--radius);
      padding: 16px;
      background: var(--panel);
      box-shadow: var(--shadow);
    }
    label {
      display: block;
      margin: 12px 0 6px;
      color: #344054;
      font-weight: 650;
      font-size: 13px;
    }
    .panel > label:first-of-type, form > label:first-of-type { margin-top: 0; }
    input, textarea, select {
      width: 100%;
      border: 1px solid var(--line);
      border-radius: 10px;
      padding: 10px 12px;
      background: #fff;
      color: var(--text);
      font: inherit;
      outline: none;
      transition: border-color .15s, box-shadow .15s;
    }
    textarea { min-height: 92px; resize: vertical; }
    input:focus, textarea:focus, select:focus {
      border-color: var(--brand);
      box-shadow: 0 0 0 3px rgba(13,148,136,.16);
    }
    input:disabled, textarea:disabled { background: #f4f6fa; color: var(--muted); cursor: not-allowed; }
    .row { display: grid; grid-template-columns: 1fr 1fr; gap: 10px; }
    .actions { display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 12px; }
    button {
      width: 100%;
      border: 0;
      border-radius: 10px;
      padding: 11px 14px;
      background: linear-gradient(135deg, var(--brand), var(--brand-dark));
      color: #fff;
      font-weight: 700;
      cursor: pointer;
      transition: transform .08s ease, box-shadow .15s, filter .15s;
      box-shadow: 0 1px 2px rgba(13,148,136,.25);
    }
    button:hover { filter: brightness(1.06); box-shadow: 0 8px 18px -7px rgba(13,148,136,.6); }
    button:active { transform: translateY(1px); }
    button.secondary { background: linear-gradient(135deg, #475569, #334155); box-shadow: none; }
    button.secondary:hover { filter: brightness(1.08); }
    button.light { background: #eef2f7; color: #2b3648; box-shadow: none; }
    button.light:hover { background: #e2e8f1; filter: none; }
    button:disabled { opacity: .6; cursor: not-allowed; transform: none; filter: none; box-shadow: none; }
    #send { margin-top: 14px; }
    #train-file { margin-top: 12px; }
    .hint {
      border: 1px solid #f3d79a;
      background: var(--warn);
      color: #7a4b00;
      border-radius: 10px;
      padding: 10px 12px;
      font-size: 12.5px;
      line-height: 1.45;
    }
    .hint.ok { border-color: #a7e6cf; background: #eafaf2; color: #067647; }
    .toolbar {
      display: flex;
      align-items: center;
      gap: 14px;
      margin-bottom: 18px;
    }
    .toolbar strong { font-size: 17px; }
    .status {
      color: var(--muted);
      font-size: 13px;
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 999px;
      padding: 4px 12px;
      font-weight: 650;
    }
    .empty {
      min-height: calc(100vh - 120px);
      display: grid;
      place-items: center;
      align-content: center;
      gap: 10px;
      color: var(--muted);
      text-align: center;
      border: 2px dashed #cdd5e4;
      border-radius: var(--radius);
      background: rgba(255,255,255,.5);
      padding: 40px;
    }
    .empty-art { font-size: 54px; line-height: 1; }
    .summary {
      background: linear-gradient(135deg, #ffffff, var(--soft));
      border: 1px solid #cdeae6;
      border-left: 5px solid var(--brand);
      border-radius: var(--radius);
      padding: 18px 20px;
      margin-bottom: 18px;
      box-shadow: var(--shadow);
    }
    .summary h2 { margin: 0 0 6px; font-size: 19px; color: var(--brand-dark); }
    .cases {
      display: grid;
      grid-template-columns: repeat(auto-fill, minmax(340px, 1fr));
      gap: 16px;
    }
    .case {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: var(--radius);
      padding: 18px;
      box-shadow: var(--shadow);
      transition: transform .12s, box-shadow .12s, border-color .12s;
    }
    .case:hover { transform: translateY(-3px); box-shadow: var(--shadow-lg); border-color: #cfe9e5; }
    .case h3 { margin: 0 0 12px; font-size: 15px; line-height: 1.35; }
    .meta { display: flex; flex-wrap: wrap; gap: 8px; margin-bottom: 14px; }
    .pill {
      border-radius: 999px;
      background: #eef2f7;
      color: #344054;
      padding: 4px 11px;
      font-size: 12px;
      font-weight: 700;
      border: 1px solid #e4e9f0;
    }
    .pill.high { background: #fee4e2; color: #912018; border-color: #fccfca; }
    .pill.medium { background: #fff3d6; color: #8a5a00; border-color: #ffe2a8; }
    .pill.low { background: #dcfae6; color: #067647; border-color: #b6f0cf; }
    .kb-status {
      display: inline-flex;
      align-items: center;
      border-radius: 999px;
      padding: 2px 9px;
      font-size: 10.5px;
      font-weight: 800;
      letter-spacing: .03em;
      margin-left: 6px;
      vertical-align: middle;
    }
    .kb-status.ready { background: #dcfae6; color: #067647; }
    .kb-status.needs_review { background: #fff1c2; color: #7a4b00; }
    .kb-status.failed { background: #fee4e2; color: #912018; }
    .section-title {
      margin: 14px 0 5px;
      color: var(--brand-dark);
      font-weight: 800;
      font-size: 11.5px;
      text-transform: uppercase;
      letter-spacing: .05em;
    }
    ol, ul { margin: 0; padding-left: 20px; }
    li { margin: 4px 0; }
    .error {
      color: var(--danger);
      background: #fff1f0;
      border: 1px solid #fecdca;
      border-radius: var(--radius);
      padding: 14px 16px;
      font-weight: 600;
    }
    .knowledge-list {
      display: grid;
      gap: 10px;
      max-height: 260px;
      overflow: auto;
      padding-right: 2px;
    }
    .knowledge-item {
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 12px;
      padding: 12px;
      font-size: 13px;
      transition: border-color .12s, box-shadow .12s;
    }
    .knowledge-item:hover { border-color: #cfe9e5; box-shadow: var(--shadow); }
    .knowledge-item strong { display: block; margin-bottom: 4px; font-size: 13.5px; }
    .item-actions {
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 6px;
      margin-top: 10px;
    }
    .item-actions button { padding: 7px 8px; font-size: 12px; }
    .review-box { display: none; margin-top: 10px; }
    .review-box.open { display: block; }
    .review-box textarea { min-height: 120px; font-size: 13px; }
    .small { color: var(--muted); font-size: 12px; }
    @media (max-width: 900px) {
      .shell { grid-template-columns: 1fr; }
      aside { border-right: 0; border-bottom: 1px solid var(--line); }
      main { padding: 18px; }
      .row, .actions { grid-template-columns: 1fr; }
      .empty { min-height: 240px; }
    }
    /* ---------- Animations ---------- */
    @keyframes fadeInUp { from { opacity: 0; transform: translateY(14px); } to { opacity: 1; transform: none; } }
    @keyframes fadeIn { from { opacity: 0; } to { opacity: 1; } }
    @keyframes pop { from { opacity: 0; transform: scale(.94); } to { opacity: 1; transform: scale(1); } }
    @keyframes spin { to { transform: rotate(360deg); } }
    @keyframes floaty { 0%, 100% { transform: translateY(0); } 50% { transform: translateY(-7px); } }
    .hero { animation: fadeInUp .5s ease both; }
    aside .panel { animation: fadeInUp .5s ease both; }
    aside .panel:nth-of-type(1) { animation-delay: .05s; }
    aside .panel:nth-of-type(2) { animation-delay: .12s; }
    aside .panel:nth-of-type(3) { animation-delay: .19s; }
    .summary { animation: fadeInUp .45s ease both; }
    .case { animation: fadeInUp .5s ease both; }
    .knowledge-item { animation: fadeIn .35s ease both; }
    .pill, .kb-status { animation: pop .3s ease both; }
    .empty-art { animation: floaty 3s ease-in-out infinite; }
    .spinner {
      width: 40px;
      height: 40px;
      border-radius: 50%;
      border: 4px solid #d7e4e1;
      border-top-color: var(--brand);
      animation: spin .8s linear infinite;
    }
    @media (prefers-reduced-motion: reduce) {
      *, *::before, *::after { animation: none !important; transition: none !important; }
    }
  </style>
</head>
<body>
  <div class="shell">
    <aside>
      <div class="hero">
        <h1><span>🧪</span> QC Test Case Agent</h1>
        <p class="subtitle">Train bằng tài liệu, workflow, diagram — rồi sinh test case bám sát ngữ cảnh đó.</p>
      </div>
      <div class="hint" id="ai-status">AI: đang kiểm tra...</div>

      <section class="panel">
        <h2>📚 Train kiến thức</h2>
        <label for="knowledge-title">Tên tài liệu / workflow</label>
        <input id="knowledge-title" placeholder="VD: Luồng thanh toán hóa đơn" />
        <label for="knowledge-type">Loại</label>
        <select id="knowledge-type">
          <option value="workflow">Workflow</option>
          <option value="diagram">Diagram</option>
          <option value="requirement">Tài liệu yêu cầu</option>
          <option value="business-rule">Business rule</option>
        </select>
        <label for="knowledge-text">Nội dung (paste cho tài liệu/workflow)</label>
        <textarea id="knowledge-text" placeholder="Dán yêu cầu, workflow, rule hoặc mô tả quy trình vào đây."></textarea>
        <div class="actions">
          <button id="train-text" type="button">＋ Train text</button>
          <button id="refresh-knowledge" class="light" type="button">↻ Làm mới</button>
        </div>
        <label for="knowledge-file">Hoặc upload file</label>
        <input id="knowledge-file" type="file" />
        <button id="train-file" class="secondary" type="button">⬆ Train từ file</button>
        <div class="hint" id="training-hint">Dùng paste cho tài liệu/workflow. Upload cho diagram. Ảnh PNG/JPG/WebP sẽ ở trạng thái NEEDS_REVIEW cho tới khi bật OCR/vision.</div>
      </section>

      <section class="panel">
        <h2>⚙️ Sinh test case</h2>
        <form id="chat-form">
          <label for="feature">Feature / Yêu cầu</label>
          <textarea id="feature" required placeholder="Mô tả feature hoặc yêu cầu cần test"></textarea>
          <div class="row">
            <div>
              <label for="actor">Actor</label>
              <input id="actor" placeholder="Vai trò người dùng" />
            </div>
            <div>
              <label for="platform">Platform</label>
              <input id="platform" placeholder="Nền tảng / màn hình" />
            </div>
          </div>
          <label for="criteria">Tiêu chí chấp nhận</label>
          <textarea id="criteria" placeholder="Mỗi tiêu chí một dòng"></textarea>
          <button id="send" type="submit">✨ Sinh test case</button>
        </form>
      </section>

      <section class="panel">
        <h2>🗂️ Knowledge base</h2>
        <div id="knowledge-list" class="knowledge-list"></div>
      </section>
    </aside>
    <main>
      <div class="toolbar">
        <strong>📋 Test case đã sinh</strong>
        <span class="status" id="status">Sẵn sàng</span>
        <button id="export-csv" class="light" type="button" style="display:none;width:auto;margin-left:auto;padding:8px 14px;font-size:13px;">⬇ Export CSV</button>
      </div>
      <div id="output" class="empty"><div class="empty-art">🧪</div><div>Train tài liệu hoặc nhập feature, rồi bấm <b>Sinh test case</b>.</div></div>
    </main>
  </div>
  <script>
    const form = document.querySelector("#chat-form");
    const output = document.querySelector("#output");
    const statusEl = document.querySelector("#status");
    const send = document.querySelector("#send");
    const exportCsvBtn = document.querySelector("#export-csv");
    const trainText = document.querySelector("#train-text");
    const trainFile = document.querySelector("#train-file");
    const refreshKnowledge = document.querySelector("#refresh-knowledge");
    const knowledgeList = document.querySelector("#knowledge-list");
    const knowledgeType = document.querySelector("#knowledge-type");
    const knowledgeText = document.querySelector("#knowledge-text");
    const trainingHint = document.querySelector("#training-hint");

    let lastGeneratedData = null;

    function escapeHtml(value) {
      return String(value ?? "").replace(/[&<>"']/g, char => ({
        "&": "&amp;",
        "<": "&lt;",
        ">": "&gt;",
        '"': "&quot;",
        "'": "&#39;"
      })[char]);
    }

    function list(items, ordered = false) {
      if (!Array.isArray(items)) {
        if (items && typeof items === "object") {
          items = Object.entries(items).map(([key, value]) => `${key}: ${Array.isArray(value) ? value.join(", ") : value}`);
        } else {
          items = items ? [items] : [];
        }
      }
      const tag = ordered ? "ol" : "ul";
      return `<${tag}>${items.map(item => `<li>${escapeHtml(item)}</li>`).join("")}</${tag}>`;
    }

    function renderKnowledge(items) {
      if (!items.length) {
        knowledgeList.innerHTML = '<div class="small">Chưa có kiến thức nào. Hãy train tài liệu đầu tiên 👆</div>';
        return;
      }
      knowledgeList.innerHTML = items.map((item, i) => `
        <div class="knowledge-item" data-id="${escapeHtml(item.id)}" style="animation-delay:${Math.min(i * 45, 360)}ms">
          <strong>${escapeHtml(item.title)}<span class="kb-status ${String(item.status || "READY").toLowerCase()}">${escapeHtml(item.status || "READY")}</span></strong>
          <div class="small">${escapeHtml(item.type)} · ${escapeHtml(item.source)} · ${item.text_length} chars · ${item.chunk_count || 0} chunks</div>
          <div class="small">${escapeHtml(item.status_message || "")}</div>
          <div>${escapeHtml(item.preview)}</div>
          <div class="item-actions">
            <button class="light" type="button" data-action="toggle-review">✎ Sửa/Review</button>
            <button class="secondary" type="button" data-action="mark-ready">✓ READY</button>
            <button class="light" type="button" data-action="mark-review">Cần review</button>
            <button class="light" type="button" data-action="delete">🗑 Xóa</button>
          </div>
          <div class="review-box">
            <label>Nội dung dùng để sinh test case</label>
            <textarea>${escapeHtml(item.text || "")}</textarea>
            <div class="actions">
              <button type="button" data-action="save-ready">Lưu READY</button>
              <button class="light" type="button" data-action="save-review">Lưu nháp</button>
            </div>
          </div>
        </div>
      `).join("");
    }

    async function updateKnowledge(id, action, text = null) {
      statusEl.textContent = "Đang cập nhật...";
      const payload = { id, action };
      if (text !== null) payload.text = text;
      const response = await fetch("/knowledge/action", {
        method: "POST",
        headers: { "Content-Type": "application/json" },
        body: JSON.stringify(payload)
      });
      const data = await response.json();
      if (!response.ok) throw new Error(data.error || "Knowledge update failed");
      await loadKnowledge();
      statusEl.textContent = data.item ? `Đã cập nhật: ${data.item.status}` : "Đã cập nhật";
    }

    function syncTrainingMode() {
      if (knowledgeType.value === "diagram") {
        knowledgeText.disabled = true;
        trainText.disabled = true;
        knowledgeText.placeholder = "Diagram nên upload dưới dạng file. Dùng SVG/drawio/BPMN/Mermaid/PDF để đọc được; PNG/JPG/WebP sẽ ở NEEDS_REVIEW.";
        trainingHint.textContent = "Chế độ diagram: hãy upload file. Ô paste bị tắt vì diagram thường là ảnh trực quan, không phải text.";
      } else {
        knowledgeText.disabled = false;
        trainText.disabled = false;
        knowledgeText.placeholder = "Dán yêu cầu, workflow, rule hoặc mô tả quy trình vào đây.";
        trainingHint.textContent = "Dùng paste cho tài liệu/workflow. Upload cho diagram. Ảnh PNG/JPG/WebP sẽ ở NEEDS_REVIEW cho tới khi bật OCR/vision.";
      }
    }

    async function loadKnowledge() {
      const response = await fetch("/knowledge");
      const data = await response.json();
      renderKnowledge(data.items || []);
    }

    async function loadAiStatus() {
      const response = await fetch("/ai/status");
      const data = await response.json();
      const aiStatus = document.querySelector("#ai-status");
      if (data.configured) {
        aiStatus.className = "hint ok";
        aiStatus.textContent = `✅ AI: ${data.model} (qua ${data.wire_api})`;
      } else {
        aiStatus.className = "hint";
        aiStatus.textContent = "⚠️ AI: chế độ fallback. Cấu hình MAAS_API_KEY hoặc LLM_API_KEY để bật LLM RAG.";
      }
    }

    function render(data) {
      const cases = data.test_cases || [];
      const sourceRefs = data.source_refs || [];
      output.className = "";
      output.innerHTML = `
        <section class="summary">
          <h2>${escapeHtml(data.feature)}</h2>
          <div>${escapeHtml(data.actor)} · ${escapeHtml(data.platform)} · ${cases.length} test case · AI: ${escapeHtml(data.ai_mode || "fallback")}</div>
          <div class="section-title">Kiến thức đã dùng</div>
          ${sourceRefs.length ? list(sourceRefs.map(ref => `${ref.title} (${ref.type})`)) : '<div class="small">Không tìm thấy kiến thức khớp.</div>'}
          <div class="section-title">Tóm tắt ngữ cảnh</div>
          <div>${escapeHtml(data.context_summary || "Không có tóm tắt ngữ cảnh.")}</div>
        </section>
        <section class="cases">
          ${cases.map((testCase, i) => `
            <article class="case" style="animation-delay:${Math.min(i * 60, 600)}ms">
              <h3>${escapeHtml(testCase.id)} · ${escapeHtml(testCase.title)}</h3>
              <div class="meta">
                <span class="pill">${escapeHtml(testCase.type)}</span>
                <span class="pill ${String(testCase.priority).toLowerCase()}">${escapeHtml(testCase.priority)}</span>
              </div>
              <div class="section-title">Điều kiện tiên quyết</div>
              ${list(testCase.preconditions || [])}
              <div class="section-title">Dữ liệu test</div>
              ${list(testCase.test_data || [])}
              <div class="section-title">Các bước</div>
              ${list(testCase.steps || [], true)}
              <div class="section-title">Kết quả mong đợi</div>
              <p>${escapeHtml(testCase.expected_result)}</p>
              ${(testCase.references || []).length ? `<div class="section-title">Tham chiếu</div>${list(testCase.references)}` : ""}
            </article>
          `).join("")}
        </section>
      `;
    }

    async function trainFromText() {
      const payload = {
        title: document.querySelector("#knowledge-title").value,
        type: document.querySelector("#knowledge-type").value,
        text: document.querySelector("#knowledge-text").value,
        source: "pasted-text"
      };
      trainText.disabled = true;
      statusEl.textContent = "Đang train text...";
      try {
        const response = await fetch("/knowledge", {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify(payload)
        });
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Training failed");
        await loadKnowledge();
        statusEl.textContent = "Đã train ✓";
        // Xoá nội dung đã nhập sau khi train thành công.
        document.querySelector("#knowledge-title").value = "";
        knowledgeText.value = "";
      } catch (error) {
        statusEl.textContent = "Lỗi";
        output.className = "error";
        output.textContent = error.message;
      } finally {
        trainText.disabled = false;
      }
    }

    async function trainFromFile() {
      const fileInput = document.querySelector("#knowledge-file");
      if (!fileInput.files.length) {
        statusEl.textContent = "Hãy chọn file trước";
        return;
      }
      const formData = new FormData();
      formData.append("file", fileInput.files[0]);
      formData.append("title", document.querySelector("#knowledge-title").value || fileInput.files[0].name);
      formData.append("type", document.querySelector("#knowledge-type").value);
      trainFile.disabled = true;
      statusEl.textContent = "Đang train file...";
      try {
        const response = await fetch("/knowledge/upload", { method: "POST", body: formData });
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Upload failed");
        await loadKnowledge();
        statusEl.textContent = `Đã lưu: ${data.item?.status || "READY"}`;
        // Xoá cache file đã chọn để lần upload sau không dính file cũ.
        fileInput.value = "";
        document.querySelector("#knowledge-title").value = "";
      } catch (error) {
        statusEl.textContent = "Lỗi";
        output.className = "error";
        output.textContent = error.message;
      } finally {
        trainFile.disabled = false;
      }
    }

    form.addEventListener("submit", async event => {
      event.preventDefault();
      send.disabled = true;
      statusEl.textContent = "Đang sinh...";
      output.className = "empty";
      output.innerHTML = '<div class="spinner"></div><div>Đang sinh test case từ knowledge base...</div>';

      const payload = {
        feature: document.querySelector("#feature").value,
        actor: document.querySelector("#actor").value,
        platform: document.querySelector("#platform").value,
        acceptance_criteria: document.querySelector("#criteria").value.split("\\n").map(v => v.trim()).filter(Boolean)
      };

      try {
        const response = await fetch("/chat", {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify(payload)
        });
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Request failed");
        lastGeneratedData = data;
        exportCsvBtn.style.display = "inline-block";
        render(data);
        statusEl.textContent = "Xong ✓";
      } catch (error) {
        output.className = "error";
        output.textContent = error.message;
        statusEl.textContent = "Lỗi";
      } finally {
        send.disabled = false;
      }
    });

    exportCsvBtn.addEventListener("click", async () => {
      if (!lastGeneratedData) return;
      try {
        exportCsvBtn.disabled = true;
        statusEl.textContent = "Đang export...";
        const response = await fetch("/export/csv", {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify({ test_cases: lastGeneratedData.test_cases || [], feature: lastGeneratedData.feature || "" })
        });
        if (!response.ok) throw new Error("Export failed");
        const blob = await response.blob();
        const url = URL.createObjectURL(blob);
        const a = document.createElement("a");
        a.href = url;
        a.download = response.headers.get("Content-Disposition")?.match(/filename="([^"]+)"/)?.[1] || "testcases.csv";
        a.click();
        URL.revokeObjectURL(url);
        statusEl.textContent = "Đã export ✓";
      } catch (error) {
        statusEl.textContent = "Lỗi export";
      } finally {
        exportCsvBtn.disabled = false;
      }
    });

    trainText.addEventListener("click", trainFromText);
    trainFile.addEventListener("click", trainFromFile);
    refreshKnowledge.addEventListener("click", loadKnowledge);
    knowledgeType.addEventListener("change", syncTrainingMode);
    knowledgeList.addEventListener("click", async event => {
      const button = event.target.closest("button[data-action]");
      if (!button) return;

      const itemEl = button.closest(".knowledge-item");
      const reviewBox = itemEl.querySelector(".review-box");
      const textarea = reviewBox.querySelector("textarea");
      const id = itemEl.dataset.id;
      const action = button.dataset.action;

      if (action === "toggle-review") {
        reviewBox.classList.toggle("open");
        return;
      }

      if (action === "delete" && !confirm("Xóa knowledge này?")) {
        return;
      }

      try {
        if (action === "mark-ready") await updateKnowledge(id, "mark-ready");
        if (action === "mark-review") await updateKnowledge(id, "mark-review");
        if (action === "delete") await updateKnowledge(id, "delete");
        if (action === "save-ready") await updateKnowledge(id, "save-ready", textarea.value);
        if (action === "save-review") await updateKnowledge(id, "save-review", textarea.value);
      } catch (error) {
        statusEl.textContent = "Lỗi";
        output.className = "error";
        output.textContent = error.message;
      }
    });
    syncTrainingMode();
    loadAiStatus();
    loadKnowledge();
  </script>
</body>
</html>
"""


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\x00", " ").strip()
    return re.sub(r"\s+", " ", text)


def _normalize_multiline(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_acceptance_criteria(raw: Any) -> list[str]:
    if isinstance(raw, list):
        items = [_clean_text(item) for item in raw]
        return [item for item in items if item]

    text = _normalize_multiline(raw)
    if not text:
        return []

    parts = re.split(r"(?:\n|;|\d+[.)]\s+|-\s+)", text)
    return [part.strip() for part in parts if part.strip()]


def _slug(value: str) -> str:
    value = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
    return value[:40] or "feature"


def _case(
    case_id: str,
    title: str,
    case_type: str,
    priority: str,
    steps: list[str],
    expected: str,
    references: list[str] | None = None,
) -> dict:
    return {
        "id": case_id,
        "title": title,
        "type": case_type,
        "priority": priority,
        "preconditions": [
            "Test environment is available and stable.",
            "Tester has an account with the required permission for this scenario.",
        ],
        "test_data": [
            "Use valid baseline data for the feature under test.",
            "Record exact input values used during execution.",
        ],
        "steps": steps,
        "expected_result": expected,
        "references": references or [],
    }


def _knowledge_lock() -> FileLock:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    lock_path = DATA_DIR / "knowledge.lock"
    return FileLock(str(lock_path))


def load_test_case_types() -> list[dict]:
    """Load custom test case types from config file, falling back to defaults."""
    if not TESTCASE_TYPES_FILE.exists():
        return list(DEFAULT_TEST_CASE_TYPES)
    try:
        data = json.loads(TESTCASE_TYPES_FILE.read_text(encoding="utf-8"))
        if isinstance(data, list) and data:
            return data
    except (json.JSONDecodeError, OSError):
        pass
    return list(DEFAULT_TEST_CASE_TYPES)


def _read_knowledge_unlocked() -> list[dict]:
    if not KNOWLEDGE_FILE.exists():
        return []
    try:
        data = json.loads(KNOWLEDGE_FILE.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Knowledge file is corrupted: {KNOWLEDGE_FILE}") from exc
    except OSError as exc:
        raise RuntimeError(f"Could not read knowledge file: {KNOWLEDGE_FILE}") from exc
    if not isinstance(data, list):
        raise RuntimeError(f"Knowledge file must contain a JSON list: {KNOWLEDGE_FILE}")
    return [item for item in data if isinstance(item, dict)]


def _read_knowledge() -> list[dict]:
    with _knowledge_lock():
        return _read_knowledge_unlocked()


def _write_knowledge_unlocked(items: list[dict]) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile("w", encoding="utf-8", dir=DATA_DIR, delete=False) as temp_file:
        json.dump(items, temp_file, indent=2, ensure_ascii=False)
        temp_file.write("\n")
        temp_path = Path(temp_file.name)
    temp_path.replace(KNOWLEDGE_FILE)


def _write_knowledge(items: list[dict]) -> None:
    with _knowledge_lock():
        _write_knowledge_unlocked(items)


def _preview(text: str, limit: int = 180) -> str:
    clean = _clean_text(text)
    if len(clean) <= limit:
        return clean
    return clean[: limit - 3] + "..."


def _keywords(text: str) -> set[str]:
    tokens = re.findall(r"(?u)\b[\w-]{3,}\b", text.lower())
    stopwords = {
        "and",
        "the",
        "for",
        "with",
        "that",
        "this",
        "from",
        "user",
        "case",
        "test",
        "feature",
        "requirement",
        "workflow",
        "diagram",
        "can",
        "after",
        "before",
        "when",
        "then",
        "must",
        "should",
        "và",
        "hoặc",
        "các",
        "cho",
        "khi",
        "sau",
        "trước",
        "người",
        "dùng",
        "tính",
        "năng",
        "kiểm",
        "thử",
    }
    return {token for token in tokens if token not in stopwords}


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[dict]:
    normalized = _normalize_multiline(text)
    if not normalized:
        return []

    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", normalized) if part.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        if len(paragraph) > chunk_size:
            if current:
                chunks.append(current.strip())
                current = ""
            start = 0
            while start < len(paragraph):
                chunks.append(paragraph[start : start + chunk_size].strip())
                start += max(1, chunk_size - overlap)
            continue

        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            chunks.append(current.strip())
            current = paragraph

    if current:
        chunks.append(current.strip())

    return [
        {
            "chunk_id": index,
            "text": chunk,
            "keywords": sorted(_keywords(chunk)),
            "preview": _preview(chunk, 160),
        }
        for index, chunk in enumerate(chunks)
        if chunk
    ]


def _chunk_score(query: str, query_keywords: set[str], item: dict, chunk: dict) -> float:
    chunk_text = " ".join(
        [
            str(item.get("title", "")),
            str(item.get("type", "")),
            str(item.get("source", "")),
            str(chunk.get("text", "")),
        ]
    )
    chunk_keywords = set(chunk.get("keywords") or _keywords(chunk_text))
    overlap = query_keywords & chunk_keywords
    # Ratio-based scoring: rewards focused chunks with high keyword density
    score: float = (len(overlap) / max(len(chunk_keywords), 1)) * 10
    lower_query = query.lower().strip()
    lower_text = chunk_text.lower()
    if lower_query and lower_query in lower_text:
        score += 10
    for term in query_keywords:
        if term in str(item.get("title", "")).lower():
            score += 2
        if term in str(item.get("type", "")).lower():
            score += 1
    return score


def add_knowledge(
    title: str,
    knowledge_type: str,
    text: str,
    source: str,
    status: str = READY,
    status_message: str = "Readable and indexed.",
    readable: bool = True,
) -> dict:
    normalized = _normalize_multiline(text)
    if status == READY and len(normalized) < 3:
        raise ValueError("Knowledge content is too short to train.")

    with _knowledge_lock():
        items = _read_knowledge_unlocked()
        item = {
            "id": f"KB-{datetime.now().strftime('%Y%m%d%H%M%S%f')}",
            "title": _clean_text(title) or "Untitled knowledge",
            "type": _clean_text(knowledge_type) or "document",
            "source": _clean_text(source) or "manual",
            "text": normalized,
            "text_length": len(normalized),
            "chunks": _chunk_text(normalized) if readable and status == READY else [],
            "preview": _preview(normalized),
            "status": status,
            "status_message": status_message,
            "readable": readable,
            "created_at": datetime.now().isoformat(),
        }
        items.insert(0, item)
        _write_knowledge_unlocked(items[:100])

    # Persist durable, shared copy to the Memory Service (best-effort) once the item is
    # readable and indexed, so trained knowledge survives restarts and is shared across replicas.
    if status == READY and readable:
        _memory_remember([f"{item['title']}\n{normalized[:8000]}"])
    return item


def update_knowledge_item(knowledge_id: str, action: str, text: str | None = None) -> dict | None:
    became_ready = False
    result: dict | None = None
    with _knowledge_lock():
        items = _read_knowledge_unlocked()
        index = next((idx for idx, item in enumerate(items) if item.get("id") == knowledge_id), None)
        if index is None:
            raise ValueError("Knowledge item not found.")

        if action == "delete":
            removed = items.pop(index)
            _write_knowledge_unlocked(items)
            return {
                "id": removed.get("id"),
                "title": removed.get("title"),
                "status": "DELETED",
            }

        item = items[index]
        now = datetime.now().isoformat()

        if action in {"save-ready", "save-review"}:
            normalized = _normalize_multiline(text)
            if action == "save-ready" and len(normalized) < 8:
                raise ValueError("Reviewed content is too short to mark READY.")
            item["text"] = normalized
            item["text_length"] = len(normalized)
            item["chunks"] = _chunk_text(normalized) if action == "save-ready" else []
            item["preview"] = _preview(normalized)
            item["reviewed_at"] = now

        if action in {"mark-ready", "save-ready"}:
            if len(_normalize_multiline(item.get("text", ""))) < 8:
                raise ValueError("Add reviewed content before marking READY.")
            item["status"] = READY
            item["readable"] = True
            item["chunks"] = _chunk_text(item.get("text", ""))
            item["status_message"] = "Reviewed and indexed."
            item["reviewed_at"] = now
            became_ready = True
        elif action in {"mark-review", "save-review"}:
            item["status"] = NEEDS_REVIEW
            item["readable"] = False
            item["chunks"] = []
            item["status_message"] = "Saved for review. It will not be used for generation until marked READY."
            item["reviewed_at"] = now
        else:
            raise ValueError(f"Unsupported knowledge action: {action}")

        items[index] = item
        _write_knowledge_unlocked(items)
        result = item

    # Persist to the Memory Service outside the lock when a reviewed item becomes READY.
    if became_ready and result:
        _memory_remember([f"{result.get('title', '')}\n{_normalize_multiline(result.get('text', ''))[:8000]}"])
    return result


def retrieve_context(query: str, limit: int = 4) -> list[dict]:
    _hydrate_from_memory()
    query_keywords = _keywords(query)
    scored = []
    for item in _read_knowledge():
        if item.get("status", READY) != READY or item.get("readable", True) is False:
            continue
        chunks = item.get("chunks") or _chunk_text(item.get("text", ""))
        for chunk in chunks:
            score = _chunk_score(query, query_keywords, item, chunk)
            if score >= 4:
                scored.append((score, item, chunk))

    scored.sort(key=lambda pair: (pair[0], pair[1].get("created_at", "")), reverse=True)

    grouped: dict[str, dict] = {}
    for score, item, chunk in scored[:TOP_CONTEXT_CHUNKS]:
        item_id = str(item.get("id"))
        if item_id not in grouped:
            grouped[item_id] = {
                **item,
                "matched_chunks": [],
                "match_score": 0,
            }
        grouped[item_id]["matched_chunks"].append(
            {
                "chunk_id": chunk.get("chunk_id"),
                "text": chunk.get("text", ""),
                "preview": chunk.get("preview") or _preview(chunk.get("text", "")),
                "score": score,
            }
        )
        grouped[item_id]["match_score"] += score

    results = sorted(grouped.values(), key=lambda item: item.get("match_score", 0), reverse=True)
    results = results[:limit]

    # Augment with durable, shared knowledge from the Memory Service (best-effort).
    if _memory_enabled():
        seen_titles = {r.get("title") for r in results}
        for mem_item in _memory_recall(query, limit=limit):
            if mem_item.get("title") not in seen_titles:
                results.append(mem_item)
                seen_titles.add(mem_item.get("title"))
        results = results[:limit]

    return results


def summarize_context(items: list[dict]) -> str:
    if not items:
        return "No matching trained document, workflow, or diagram was found."

    snippets = []
    for item in items:
        matched_text = " ".join(chunk.get("text", "") for chunk in item.get("matched_chunks", []))
        snippets.append(f"{item.get('title', 'Untitled')}: {_preview(matched_text or item.get('text', ''), 360)}")
    summary = " | ".join(snippets)
    return summary[:MAX_CONTEXT_CHARS]


def derive_criteria_from_context(items: list[dict]) -> list[str]:
    candidates: list[str] = []
    for item in items:
        text = _normalize_multiline(
            "\n".join(chunk.get("text", "") for chunk in item.get("matched_chunks", [])) or item.get("text", "")
        )
        lines = [line.strip(" -\t") for line in text.splitlines() if line.strip()]
        for line in lines:
            lower = line.lower()
            if any(word in lower for word in ["must", "should", "cannot", "reject", "approve", "expire", "timeout", "permission", "role", "rule"]):
                candidates.append(line)

    cleaned = []
    seen = set()
    for candidate in candidates:
        value = _clean_text(candidate).strip(":- ")
        lower_value = value.lower()
        if lower_value in {"rule", "rules", "business rule", "business rules", "acceptance criteria"}:
            continue
        if 8 <= len(value) <= 180 and lower_value not in seen:
            seen.add(value.lower())
            cleaned.append(value)
    return cleaned[:8]


def derive_workflow_steps(items: list[dict]) -> list[str]:
    for item in items:
        text = _normalize_multiline(
            "\n".join(chunk.get("text", "") for chunk in item.get("matched_chunks", [])) or item.get("text", "")
        )
        if "->" in text:
            first_flow = next((line for line in text.splitlines() if "->" in line), "")
            steps = [_clean_text(part) for part in first_flow.split("->")]
            steps = [step for step in steps if step]
            if len(steps) >= 2:
                return steps[:10]
    return []


def llm_status() -> dict:
    api_key = os.getenv("LLM_API_KEY") or os.getenv("MAAS_API_KEY")
    base_url = os.getenv("LLM_BASE_URL") or "https://maas-llm-aiplatform-hcm.api.vngcloud.vn/v1"
    model = os.getenv("LLM_MODEL") or "minimax/minimax-m2.5"
    wire_api = os.getenv("LLM_WIRE_API") or ("responses" if os.getenv("MAAS_API_KEY") and not os.getenv("LLM_WIRE_API") else "chat")
    configured = all([api_key, base_url, model])
    return {
        "configured": configured,
        "base_url": base_url if configured else "",
        "model": model if configured else "",
        "wire_api": wire_api,
        "env_key": "LLM_API_KEY" if os.getenv("LLM_API_KEY") else "MAAS_API_KEY",
    }


def _context_for_llm(items: list[dict]) -> str:
    blocks = []
    for item in items:
        chunks = item.get("matched_chunks") or [{"text": item.get("text", "")}]
        for chunk in chunks:
            text = _normalize_multiline(chunk.get("text", ""))
            if text:
                blocks.append(
                    f"Source: {item.get('title', 'Untitled')} | Type: {item.get('type', 'document')} | "
                    f"Chunk: {chunk.get('chunk_id', 'full')}\n{text}"
                )
    return "\n\n---\n\n".join(blocks)[:MAX_CONTEXT_CHARS]


def _coerce_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [_clean_text(item) for item in value if _clean_text(item)]
    if isinstance(value, dict):
        items = []
        for key, item_value in value.items():
            if isinstance(item_value, (list, dict)):
                item_text = json.dumps(item_value, ensure_ascii=False)
            else:
                item_text = _clean_text(item_value)
            key_text = _clean_text(key)
            if key_text and item_text:
                items.append(f"{key_text}: {item_text}")
        return items
    text = _clean_text(value)
    return [text] if text else []


def _sanitize_test_case(raw_case: dict, fallback_index: int) -> dict:
    return {
        "id": _clean_text(raw_case.get("id")) or f"TC-LLM-{fallback_index:03d}",
        "title": _clean_text(raw_case.get("title")) or "Generated test case",
        "type": _clean_text(raw_case.get("type")) or "Functional",
        "priority": _clean_text(raw_case.get("priority")) or "Medium",
        "preconditions": _coerce_list(raw_case.get("preconditions")),
        "test_data": _coerce_list(raw_case.get("test_data")),
        "steps": _coerce_list(raw_case.get("steps")),
        "expected_result": _clean_text(raw_case.get("expected_result")) or "Expected behavior is observed.",
        "references": _coerce_list(raw_case.get("references")),
    }


def _parse_llm_json(content: str) -> dict:
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content, flags=re.IGNORECASE | re.DOTALL).strip()

    decoder = json.JSONDecoder()
    candidates = [content]

    first_brace = content.find("{")
    last_brace = content.rfind("}")
    if first_brace >= 0 and last_brace > first_brace:
        candidates.append(content[first_brace : last_brace + 1])
        candidates.append(content[first_brace:])

    last_error: Exception | None = None
    for candidate in candidates:
        try:
            parsed = json.loads(candidate)
        except json.JSONDecodeError as exc:
            last_error = exc
            try:
                parsed, _ = decoder.raw_decode(candidate)
            except json.JSONDecodeError as raw_exc:
                last_error = raw_exc
                continue
        if isinstance(parsed, dict):
            return parsed

    raise ValueError(f"Could not parse LLM JSON response: {last_error}")


# --- Persistent long-term memory (AgentBase Memory Service) ---
# Stores auto-learned QC knowledge durably and shared across replicas, replacing the
# ephemeral local knowledge.json. All calls are best-effort: any failure falls back to
# local-only behavior so the agent never breaks.
MEMORY_RECORDS_API = "https://agentbase.api.vngcloud.vn/memory/memories"
_memory_client = None


def _memory_enabled() -> bool:
    return bool(os.getenv("MEMORY_ID") and os.getenv("MEMORY_STRATEGY_ID"))


def _memory_namespace() -> str:
    strategy = os.getenv("MEMORY_STRATEGY_ID", "")
    actor = os.getenv("MEMORY_ACTOR", "qc-shared")
    return f"/strategies/{strategy}/actors/{actor}"


def _memory_token() -> str | None:
    global _memory_client
    try:
        if _memory_client is None:
            from greennode_agentbase.memory import MemoryClient

            _memory_client = MemoryClient()
        return _memory_client._get_oauth2_token_sync()
    except Exception:
        return None


def _memory_remember(facts: list[str]) -> bool:
    if not _memory_enabled():
        return False
    facts = [f.strip() for f in facts if f and f.strip()]
    if not facts:
        return False
    token = _memory_token()
    if not token:
        return False
    try:
        import httpx

        mid = os.getenv("MEMORY_ID")
        resp = httpx.post(
            f"{MEMORY_RECORDS_API}/{mid}/memory-records:insert-directly",
            params={"namespace": _memory_namespace()},
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            json={"memoryRecords": facts},
            timeout=20,
        )
        resp.raise_for_status()
        return True
    except Exception:
        return False


def _memory_recall(query: str, limit: int = 4) -> list[dict]:
    """Semantic search persistent memory; return items shaped like local knowledge context items."""
    if not _memory_enabled() or not query.strip():
        return []
    token = _memory_token()
    if not token:
        return []
    try:
        import httpx

        mid = os.getenv("MEMORY_ID")
        resp = httpx.post(
            f"{MEMORY_RECORDS_API}/{mid}/memory-records:search",
            params={"namespace": _memory_namespace()},
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            json={"query": query, "limit": max(5, limit)},
            timeout=20,
        )
        resp.raise_for_status()
        records = resp.json()
        if not isinstance(records, list):
            records = records.get("listData") or records.get("content") or []
        items = []
        for rec in records[:limit]:
            text = _normalize_multiline(rec.get("memory", ""))
            if not text:
                continue
            items.append(
                {
                    "id": f"MEM-{rec.get('id')}",
                    "title": (text.split(".")[0] or "Memory")[:80],
                    "type": AUTO_LEARNED_TYPE,
                    "source": "memory-service",
                    "text": text,
                    "matched_chunks": [
                        {
                            "chunk_id": "mem",
                            "text": text,
                            "preview": _preview(text),
                            "score": rec.get("score"),
                        }
                    ],
                    "match_score": rec.get("score") or 0,
                }
            )
        return items
    except Exception:
        return []


_memory_hydrated = False


def _memory_list(limit: int = 100) -> list[dict]:
    """List every persisted memory record for this namespace. Used to rebuild the
    local knowledge store after a restart wipes the ephemeral .agentbase volume."""
    if not _memory_enabled():
        return []
    token = _memory_token()
    if not token:
        return []
    try:
        import httpx

        mid = os.getenv("MEMORY_ID")
        resp = httpx.get(
            f"{MEMORY_RECORDS_API}/{mid}/memory-records",
            params={"namespace": _memory_namespace(), "limit": limit},
            headers={"Authorization": f"Bearer {token}"},
            timeout=20,
        )
        resp.raise_for_status()
        records = resp.json()
        if not isinstance(records, list):
            records = records.get("listData") or records.get("content") or []
        return [rec for rec in records if isinstance(rec, dict)]
    except Exception:
        return []


def _hydrate_from_memory() -> None:
    """Restore trained knowledge from the durable Memory Service into the local store
    on first access. The runtime's local .agentbase volume is ephemeral, so after a
    restart the local knowledge.json is empty even though the Memory Service still holds
    every READY item — this is why trained knowledge appeared to be "lost" on the UI.
    Runs at most once per process and is fully best-effort: any failure leaves the
    request path untouched."""
    global _memory_hydrated
    if _memory_hydrated or not _memory_enabled():
        return
    _memory_hydrated = True  # mark first so a failure never re-runs on every request
    records = _memory_list(limit=100)
    if not records:
        return
    try:
        with _knowledge_lock():
            items = _read_knowledge_unlocked()
            seen = {_normalize_multiline(it.get("text", "")) for it in items}
            restored = []
            for rec in records:
                raw = _normalize_multiline(rec.get("memory", ""))
                if not raw:
                    continue
                title, _, body = raw.partition("\n")
                body = body.strip() or raw
                if body in seen:
                    continue
                seen.add(body)
                restored.append(
                    {
                        "id": f"MEM-{rec.get('id')}",
                        "title": _clean_text(title)[:120] or "Restored memory",
                        "type": AUTO_LEARNED_TYPE,
                        "source": "memory-service",
                        "text": body,
                        "text_length": len(body),
                        "chunks": _chunk_text(body),
                        "preview": _preview(body),
                        "status": READY,
                        "status_message": "Restored from Memory Service.",
                        "readable": True,
                        "created_at": rec.get("created_at") or datetime.now().isoformat(),
                    }
                )
            if restored:
                _write_knowledge_unlocked((items + restored)[:100])
    except Exception:
        return


def _auto_learn(feature: str, learned_summary: str) -> dict | None:
    """Persist the model's distilled understanding of a feature so it can be reused
    as business context next time. Only runs after a self-generated (no-context) run."""
    feature = _clean_text(feature)
    summary = _normalize_multiline(learned_summary)
    if not feature or feature == "Unspecified feature" or len(summary) < 40:
        return None

    title = f"Auto-learned: {feature}"
    # Dedup: do not create a second auto-learned record for the same feature.
    for item in _read_knowledge():
        if item.get("type") == AUTO_LEARNED_TYPE and item.get("title") == title:
            return None

    # add_knowledge() persists the READY item to the Memory Service (best-effort), so it
    # survives restarts and is shared across replicas. retrieve_context already confirmed no
    # match existed for this feature before self-generation, so this will not create duplicates.
    try:
        return add_knowledge(
            title=title,
            knowledge_type=AUTO_LEARNED_TYPE,
            text=summary,
            source="self-learn",
            status=READY,
            status_message="Auto-learned from self-generated test cases.",
            readable=True,
        )
    except Exception:
        return None


def _enhance_with_llm(response: dict, payload: dict, context_items: list[dict]) -> dict:
    status = llm_status()
    if not status["configured"]:
        response["ai_mode"] = "fallback-no-llm-config"
        response["ai_status"] = status
        return response

    try:
        import httpx

        base_url = status["base_url"].rstrip("/")
        api_key = os.getenv("LLM_API_KEY") or os.getenv("MAAS_API_KEY", "")
        model = status["model"]
        wire_api = status["wire_api"]
        context = _context_for_llm(context_items)
        has_context = bool(context_items) and bool(context.strip())
        fallback_cases = response.get("test_cases", [])
        try:
            timeout_s = float(os.getenv("LLM_TIMEOUT", "120") or 120)
        except ValueError:
            timeout_s = 120.0

        if has_context:
            instructions = (
                "You are a senior QC test design assistant. Use the filtered business context to improve "
                "and ground the provided baseline test cases. Prefer facts from the context and do not "
                "invent business rules that contradict it. Return ONLY valid JSON with keys test_cases "
                "and notes."
            )
            context_block = context
        else:
            instructions = (
                "You are a senior QC test design assistant. No trained business context is available for "
                "this request, so generate the test cases yourself from your own QC expertise. Design "
                "comprehensive, realistic test cases for the feature using standard QC best practices — "
                "cover positive, negative, boundary, permission/security, resilience, and any "
                "feature-specific scenarios. Use the baseline test cases only as a starting point and "
                "expand well beyond them. Clearly state any assumptions you make in the notes. Return "
                "ONLY valid JSON with keys test_cases, notes, and learned_summary. The learned_summary "
                "is a concise 3-8 sentence distilled understanding of this feature (purpose, key actors, "
                "main flows, business rules, and assumptions used) written so it can be reused as "
                "business context next time."
            )
            context_block = "No trained context. Generate from your own QC expertise."

        feature = _clean_text(response.get("feature") or payload.get("feature")) or "Unspecified feature"
        actor = _clean_text(response.get("actor") or payload.get("actor") or "User")
        platform = _clean_text(response.get("platform") or payload.get("platform") or "Target application")
        criteria = _split_acceptance_criteria(payload.get("acceptance_criteria") or payload.get("criteria"))
        criteria_block = "\n".join(f"- {c}" for c in criteria) if criteria else "None provided."
        baseline_json = json.dumps(fallback_cases, ensure_ascii=False, indent=2)
        learned_line = ',\n  "learned_summary": "<3-8 sentence reusable summary>"' if not has_context else ""
        prompt_text = (
            f"{instructions}\n\n"
            f"Feature: {feature}\n"
            f"Actor: {actor}\n"
            f"Platform: {platform}\n"
            f"Acceptance criteria:\n{criteria_block}\n\n"
            f"Business context:\n{context_block}\n\n"
            "Baseline rule-based test cases (JSON array) to use as a starting point and expand well beyond:\n"
            f"{baseline_json}\n\n"
            "Return ONLY a single valid JSON object (no markdown fences) with this exact schema:\n"
            "{\n"
            '  "test_cases": [\n'
            '    {"id": "string", "title": "string", "type": "string", '
            '"priority": "High|Medium|Low", "preconditions": ["string"], '
            '"test_data": ["string"], "steps": ["string"], '
            '"expected_result": "string", "references": ["string"]}\n'
            "  ],\n"
            '  "notes": ["string"]'
            f"{learned_line}\n"
            "}"
        )

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        if wire_api == "responses":
            llm_response = httpx.post(
                f"{base_url}/responses",
                headers=headers,
                json={
                    "model": model,
                    "input": [
                        {
                            "role": "system",
                            "content": "Return compact, valid JSON only. No markdown fences.",
                        },
                        {
                            "role": "user",
                            "content": prompt_text,
                        },
                    ],
                    "temperature": 0.2,
                    "max_output_tokens": 8000,
                },
                timeout=timeout_s,
            )
        else:
            llm_response = httpx.post(
                f"{base_url}/chat/completions",
                headers=headers,
                json={
                    "model": model,
                    "messages": [
                        {
                            "role": "system",
                            "content": "Return compact, valid JSON only. No markdown fences.",
                        },
                        {
                            "role": "user",
                            "content": prompt_text,
                        },
                    ],
                    "temperature": 0.2,
                },
                timeout=timeout_s,
            )
        llm_response.raise_for_status()
        llm_payload = llm_response.json()
        if wire_api == "responses":
            content = (
                llm_payload.get("output_text")
                or "".join(
                    part.get("text", "")
                    for item in llm_payload.get("output", [])
                    for part in item.get("content", [])
                    if part.get("type") in {"output_text", "text"}
                )
            )
        else:
            content = llm_payload["choices"][0]["message"]["content"]

        parsed = _parse_llm_json(content)
        if isinstance(parsed.get("test_cases"), list) and parsed["test_cases"]:
            sanitized_cases = [
                _sanitize_test_case(test_case, index)
                for index, test_case in enumerate(parsed["test_cases"], start=1)
                if isinstance(test_case, dict)
            ]
            if not sanitized_cases:
                raise ValueError("LLM response did not include valid test case objects.")
            response["test_cases"] = sanitized_cases
            response["notes"] = _coerce_list(parsed.get("notes")) or response.get("notes", [])
            response["ai_mode"] = "llm-rag" if has_context else "llm-generate"
            response["ai_status"] = status
            if not has_context:
                learned_item = _auto_learn(response.get("feature", ""), parsed.get("learned_summary", ""))
                if learned_item:
                    response["auto_learned"] = {
                        "id": learned_item.get("id"),
                        "title": learned_item.get("title"),
                    }
            return response
        raise ValueError("LLM response did not include test_cases.")
    except Exception as exc:
        response["ai_mode"] = "fallback-llm-error"
        response["ai_error"] = str(exc)
        response["ai_status"] = status
        return response


def _extract_image_via_vision(data: bytes, media_type: str) -> str:
    """Extract text from an image using a vision-capable LLM.

    Only active when LLM_WIRE_API=chat and LLM_VISION_ENABLED=true.
    Raises ValueError on failure or when vision is not configured.
    """
    status = llm_status()
    vision_enabled = os.getenv("LLM_VISION_ENABLED", "").lower() == "true"
    if not status["configured"] or status["wire_api"] != "chat" or not vision_enabled:
        raise ValueError("Vision extraction not available: check LLM_WIRE_API=chat and LLM_VISION_ENABLED=true.")

    import httpx

    base_url = status["base_url"].rstrip("/")
    api_key = os.getenv("LLM_API_KEY") or os.getenv("MAAS_API_KEY", "")
    model = status["model"]
    image_b64 = base64.b64encode(data).decode("ascii")

    response = httpx.post(
        f"{base_url}/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={
            "model": model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{media_type};base64,{image_b64}"},
                        },
                        {
                            "type": "text",
                            "text": (
                                "Extract all text and structured content from this image for QA documentation. "
                                "Include labels, states, transitions, steps, conditions, and any visible text. "
                                "Format the output as readable plain text."
                            ),
                        },
                    ],
                }
            ],
            "temperature": 0.1,
        },
        timeout=30,
    )
    response.raise_for_status()
    content = response.json()["choices"][0]["message"]["content"]
    if not content or not content.strip():
        raise ValueError("Vision model returned empty response.")
    return content.strip()


def build_test_cases(payload: dict) -> dict:
    feature = _clean_text(payload.get("feature") or payload.get("message") or payload.get("requirement"))
    if not feature:
        feature = "Unspecified feature"

    actor = _clean_text(payload.get("actor") or payload.get("user_role") or "User")
    platform = _clean_text(payload.get("platform") or "Target application")
    criteria = _split_acceptance_criteria(payload.get("acceptance_criteria") or payload.get("criteria"))

    query = " ".join([feature, actor, platform, " ".join(criteria)])
    context_items = retrieve_context(query)
    context_summary = summarize_context(context_items)
    context_refs = [f"{item.get('title', 'Untitled')} ({item.get('type', 'document')})" for item in context_items]

    if not criteria:
        criteria = derive_criteria_from_context(context_items)

    workflow_steps = derive_workflow_steps(context_items)
    feature_key = _slug(feature)

    type_steps_map: dict[str, tuple[list[str], str]] = {
        "Positive": (
            [
                f"Open {platform} with a valid {actor} account.",
                f"Navigate to the {feature} flow.",
                "Prepare data according to the trained business document.",
                "Submit or complete the action.",
            ],
            f"The {feature} action completes successfully and the system matches the documented workflow.",
        ),
        "Negative": (
            [
                f"Open the {feature} flow.",
                "Leave mandatory fields empty or provide invalid values from the trained rules.",
                "Submit the form or trigger the action.",
            ],
            "The system blocks submission and shows clear validation messages near the invalid fields.",
        ),
        "Boundary": (
            [
                f"Open the {feature} flow.",
                "Enter minimum allowed values and submit.",
                "Repeat with maximum allowed values.",
                "Repeat with values just outside the allowed range.",
            ],
            "Allowed boundary values are accepted; out-of-range values are rejected with understandable errors.",
        ),
        "Permission": (
            [
                "Sign in with an account that should not have access to this feature.",
                f"Attempt to open or execute the {feature} flow.",
            ],
            "The system denies access without exposing sensitive data or allowing the action to complete.",
        ),
        "Resilience": (
            [
                f"Start the {feature} flow with valid data.",
                "Simulate a timeout, refresh, duplicate submit, or network interruption.",
                "Return to the feature and verify the final state.",
            ],
            "The system prevents duplicate/partial inconsistent results and gives the user a recoverable state.",
        ),
        "Workflow": (
            [f"Verify workflow step: {step}." for step in workflow_steps],
            "Each transition follows the trained workflow and invalid transitions are not allowed.",
        ),
    }

    active_types = load_test_case_types()
    cases: list[dict] = []
    case_index = 1
    workflow_included = False

    for type_def in active_types:
        if not type_def.get("enabled", True):
            continue
        condition = type_def.get("condition")
        case_type = type_def.get("type", "")
        priority = type_def.get("priority", "Medium")

        if condition == "has_workflow" and not workflow_steps:
            continue

        if case_type == "Workflow":
            workflow_included = True

        steps_hint = type_def.get("steps_hint")
        if steps_hint:
            steps = [steps_hint]
            expected = f"The system satisfies the {case_type} criteria for {feature}."
        elif case_type in type_steps_map:
            steps, expected = type_steps_map[case_type]
        else:
            steps = [
                f"Open the {feature} flow.",
                f"Execute the {case_type.lower()} scenario for {feature}.",
                "Observe the system response.",
            ]
            expected = f"The system handles the {case_type.lower()} scenario correctly for {feature}."

        title_map: dict[str, str] = {
            "Positive": f"{actor} can complete the happy path for {feature}",
            "Negative": f"Required validation is shown for missing or invalid data in {feature}",
            "Boundary": f"Boundary values are handled correctly for {feature}",
            "Permission": f"Unauthorized access is prevented for {feature}",
            "Resilience": f"System handles interruption or failure during {feature}",
            "Workflow": f"Workflow transitions follow the trained diagram for {feature}",
        }
        title = title_map.get(case_type) or type_def.get("title_template", f"{case_type} scenario for {feature}")
        cases.append(_case(f"TC-{feature_key}-{case_index:03d}", title, case_type, priority, steps, expected, context_refs))
        case_index += 1

    start_index = case_index
    for index, criterion in enumerate(criteria, start=start_index):
        cases.append(
            _case(
                f"TC-{feature_key}-{index:03d}",
                f"Acceptance criterion is satisfied: {criterion}",
                "Acceptance",
                "High",
                [
                    f"Prepare data and user state for: {criterion}.",
                    f"Execute the {feature} behavior related to this criterion.",
                    "Observe system response, stored data, events, and user-visible result.",
                ],
                f"The system behavior matches the acceptance criterion: {criterion}.",
                context_refs,
            )
        )

    # Deduplicate cases with identical titles (case-insensitive)
    seen_titles: set[str] = set()
    unique_cases = []
    for case in cases:
        title_key = case.get("title", "").lower().strip()
        if title_key not in seen_titles:
            seen_titles.add(title_key)
            unique_cases.append(case)
    cases = unique_cases

    response = {
        "status": "success",
        "generated_at": datetime.now().isoformat(),
        "feature": feature,
        "actor": actor,
        "platform": platform,
        "context_summary": context_summary,
        "source_refs": [
            {
                "id": item.get("id"),
                "title": item.get("title"),
                "type": item.get("type"),
                "source": item.get("source"),
                "matched_chunks": [
                    {
                        "chunk_id": chunk.get("chunk_id"),
                        "preview": chunk.get("preview"),
                        "score": chunk.get("score"),
                    }
                    for chunk in item.get("matched_chunks", [])
                ],
            }
            for item in context_items
        ],
        "test_cases": cases,
        "quality_checklist": QUALITY_CHECKLIST,
        "notes": [
            "Upload requirement documents, workflow exports, diagrams, or business rules before generating final test cases.",
            "Image-only diagrams need OCR/vision integration; export them as SVG, drawio, Mermaid, BPMN, PDF, or paste the flow text for best local results.",
        ],
    }
    return _enhance_with_llm(response, payload, context_items)


def _decode_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _strip_markup(text: str) -> str:
    text = unescape(text)
    text = re.sub(r"<[^>]+>", " ", text)
    return _normalize_multiline(text)


def _strip_markdown(text: str) -> str:
    # Fenced code blocks: keep content, remove fences
    text = re.sub(r"^```[^\n]*\n(.*?)```", lambda m: m.group(1), text, flags=re.MULTILINE | re.DOTALL)
    # Inline code
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # ATX headers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Setext headers
    text = re.sub(r"^[=\-]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Bold/italic
    text = re.sub(r"\*{1,3}([^*\n]+)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_\n]+)_{1,3}", r"\1", text)
    # Images before links
    text = re.sub(r"!\[([^\]]*)\]\([^\)]*\)", r"\1", text)
    # Links
    text = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", text)
    # Blockquotes
    text = re.sub(r"^>\s*", "", text, flags=re.MULTILINE)
    # Horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # HTML tags
    text = re.sub(r"<[^>]+>", " ", text)
    # List markers (- item, * item, 1. item)
    text = re.sub(r"^[\s]*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[\s]*\d+\.\s+", "", text, flags=re.MULTILINE)
    return _normalize_multiline(text)


def _max_upload_bytes() -> int:
    raw = os.getenv("MAX_UPLOAD_BYTES", str(DEFAULT_MAX_UPLOAD_BYTES))
    try:
        return max(1, int(raw))
    except ValueError:
        return DEFAULT_MAX_UPLOAD_BYTES


def _format_bytes(value: int) -> str:
    if value >= 1024 * 1024:
        return f"{value / (1024 * 1024):.1f} MB"
    if value >= 1024:
        return f"{value / 1024:.1f} KB"
    return f"{value} bytes"


def _write_temp_upload(data: bytes, suffix: str) -> Path:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(delete=False, dir=DATA_DIR, suffix=suffix) as temp_file:
        temp_file.write(data)
        return Path(temp_file.name)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("python-docx is not installed. Run: venv/bin/pip install -r requirements.txt") from exc

    temp_path = _write_temp_upload(data, ".docx")
    try:
        document = Document(str(temp_path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return _normalize_multiline("\n".join(paragraphs))
    finally:
        temp_path.unlink(missing_ok=True)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("pypdf is not installed. Run: venv/bin/pip install -r requirements.txt") from exc

    temp_path = _write_temp_upload(data, ".pdf")
    try:
        reader = PdfReader(str(temp_path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return _normalize_multiline("\n".join(pages))
    finally:
        temp_path.unlink(missing_ok=True)


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ValueError("openpyxl is not installed. Run: venv/bin/pip install -r requirements.txt") from exc

    temp_path = _write_temp_upload(data, ".xlsx")
    try:
        workbook = load_workbook(str(temp_path), read_only=True, data_only=True)
        rows = []
        for sheet in workbook.worksheets:
            rows.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [_clean_text(value) for value in row if value is not None and _clean_text(value)]
                if values:
                    rows.append(" | ".join(values))
        return _normalize_multiline("\n".join(rows))
    finally:
        temp_path.unlink(missing_ok=True)


def extract_file_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")

    if suffix in IMAGE_EXTENSIONS:
        media_type_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp"}
        media_type = media_type_map.get(suffix, "image/png")
        return _extract_image_via_vision(data, media_type)

    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix == ".xlsx":
        return _extract_xlsx(data)

    text = _decode_bytes(data)
    if suffix in {".md"}:
        stripped = _strip_markdown(text)
        return stripped if stripped else _normalize_multiline(text)

    if suffix in {".xml", ".svg", ".drawio", ".bpmn"}:
        try:
            root = ElementTree.fromstring(text)
            xml_texts = [element.text.strip() for element in root.iter() if element.text and element.text.strip()]
            attributes = [str(value) for element in root.iter() for value in element.attrib.values() if value]
            text = "\n".join(xml_texts + attributes) or text
        except ElementTree.ParseError:
            pass
        return _strip_markup(text)

    if suffix == ".json":
        try:
            parsed = json.loads(text)
            return _normalize_multiline(json.dumps(parsed, indent=2, ensure_ascii=False))
        except json.JSONDecodeError:
            return _normalize_multiline(text)

    return _normalize_multiline(text)


async def chat_page(request: Request) -> HTMLResponse:
    return HTMLResponse(CHAT_HTML)


async def chat_api(request: Request) -> JSONResponse:
    try:
        payload = await request.json()
    except Exception:
        return JSONResponse({"error": "Invalid JSON payload"}, status_code=400)

    if not isinstance(payload, dict):
        return JSONResponse({"error": "Payload must be a JSON object"}, status_code=400)

    try:
        return JSONResponse(build_test_cases(payload))
    except RuntimeError as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)


async def ai_status_api(request: Request) -> JSONResponse:
    return JSONResponse(llm_status())


async def knowledge_list_api(request: Request) -> JSONResponse:
    items = []
    _hydrate_from_memory()
    try:
        knowledge_items = _read_knowledge()
    except RuntimeError as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)

    for item in knowledge_items:
        items.append(
            {
                "id": item.get("id"),
                "title": item.get("title"),
                "type": item.get("type"),
                "source": item.get("source"),
                "text": item.get("text", ""),
                "text_length": item.get("text_length", len(item.get("text", ""))),
                "chunk_count": len(item.get("chunks") or _chunk_text(item.get("text", ""))),
                "preview": item.get("preview") or _preview(item.get("text", "")),
                "status": item.get("status", READY),
                "status_message": item.get("status_message", "Readable and indexed."),
                "readable": item.get("readable", True),
                "created_at": item.get("created_at"),
            }
        )
    return JSONResponse({"items": items})


async def knowledge_action_api(request: Request) -> JSONResponse:
    try:
        payload = await request.json()
    except Exception:
        return JSONResponse({"error": "Invalid JSON payload"}, status_code=400)

    knowledge_id = _clean_text(payload.get("id"))
    action = _clean_text(payload.get("action"))
    if not knowledge_id or not action:
        return JSONResponse({"error": "Both id and action are required."}, status_code=400)

    try:
        item = update_knowledge_item(knowledge_id, action, payload.get("text"))
    except ValueError as exc:
        return JSONResponse({"error": str(exc)}, status_code=400)
    except RuntimeError as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)

    return JSONResponse({"item": item})


async def knowledge_create_api(request: Request) -> JSONResponse:
    try:
        payload = await request.json()
    except Exception:
        return JSONResponse({"error": "Invalid JSON payload"}, status_code=400)

    knowledge_type = str(payload.get("type", "document"))
    if knowledge_type == "diagram":
        try:
            item = add_knowledge(
                title=payload.get("title", "Untitled diagram"),
                knowledge_type=knowledge_type,
                text="Diagram was submitted through the text API. Upload the original diagram file so it can be parsed or marked for OCR/vision review.",
                source=payload.get("source", "pasted-text"),
                status=NEEDS_REVIEW,
                status_message="Diagram text paste is disabled. Upload SVG/drawio/BPMN/Mermaid/PDF, or connect OCR/vision for image diagrams.",
                readable=False,
            )
        except RuntimeError as exc:
            return JSONResponse({"error": str(exc)}, status_code=500)
        return JSONResponse({"item": item}, status_code=202)

    try:
        item = add_knowledge(
            title=payload.get("title", "Untitled knowledge"),
            knowledge_type=knowledge_type,
            text=payload.get("text", ""),
            source=payload.get("source", "manual"),
        )
    except ValueError as exc:
        return JSONResponse({"error": str(exc)}, status_code=400)
    except RuntimeError as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)

    return JSONResponse({"item": item})


async def knowledge_upload_api(request: Request) -> JSONResponse:
    max_upload_bytes = _max_upload_bytes()
    content_length = request.headers.get("content-length")
    if content_length:
        try:
            if int(content_length) > max_upload_bytes:
                return JSONResponse(
                    {"error": f"Upload is too large. Maximum size is {_format_bytes(max_upload_bytes)}."},
                    status_code=413,
                )
        except ValueError:
            return JSONResponse({"error": "Invalid Content-Length header."}, status_code=400)

    try:
        form = await request.form()
    except Exception as exc:
        return JSONResponse({"error": f"Could not read upload form: {exc}"}, status_code=400)

    upload = form.get("file")
    if upload is None:
        return JSONResponse({"error": "Missing file upload"}, status_code=400)

    filename = getattr(upload, "filename", "uploaded-file")
    suffix = Path(filename).suffix.lower()
    try:
        data = await upload.read()
        if not data:
            return JSONResponse({"error": "Uploaded file is empty."}, status_code=400)
        if len(data) > max_upload_bytes:
            return JSONResponse(
                {"error": f"Upload is too large. Maximum size is {_format_bytes(max_upload_bytes)}."},
                status_code=413,
            )
        if suffix in IMAGE_EXTENSIONS:
            media_type_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp"}
            media_type = media_type_map.get(suffix, "image/png")
            try:
                extracted_text = _extract_image_via_vision(data, media_type)
                item = add_knowledge(
                    title=str(form.get("title") or filename),
                    knowledge_type=str(form.get("type") or "diagram"),
                    text=extracted_text,
                    source=filename,
                    status=READY,
                    status_message="Extracted via vision model and indexed.",
                )
                return JSONResponse({"item": item})
            except Exception:
                pass
            item = add_knowledge(
                title=str(form.get("title") or filename),
                knowledge_type=str(form.get("type") or "diagram"),
                text=f"Image diagram file uploaded: {filename}. Configure LLM_VISION_ENABLED=true with a vision-capable model to extract content automatically.",
                source=filename,
                status=NEEDS_REVIEW,
                status_message="Saved, but not indexed. Export diagram as SVG/drawio/BPMN/Mermaid/PDF or set LLM_VISION_ENABLED=true.",
                readable=False,
            )
            return JSONResponse({"item": item}, status_code=202)

        text = extract_file_text(filename, data)
        item = add_knowledge(
            title=str(form.get("title") or filename),
            knowledge_type=str(form.get("type") or "document"),
            text=text,
            source=filename,
        )
    except ValueError as exc:
        item = add_knowledge(
            title=str(form.get("title") or filename),
            knowledge_type=str(form.get("type") or "document"),
            text=f"File upload could not be indexed: {filename}. Reason: {exc}",
            source=filename,
            status=FAILED,
            status_message=str(exc),
            readable=False,
        )
        return JSONResponse({"item": item}, status_code=202)
    except RuntimeError as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)
    except Exception as exc:
        return JSONResponse({"error": f"Could not train file: {exc}"}, status_code=500)

    return JSONResponse({"item": item})


async def config_types_api(request: Request) -> JSONResponse:
    return JSONResponse({"types": load_test_case_types()})


async def export_csv_api(request: Request) -> Response:
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "Invalid JSON body"}, status_code=400)

    test_cases = body.get("test_cases", [])
    feature = body.get("feature", "testcases")

    sio = io.StringIO()
    writer = csv.writer(sio)
    writer.writerow(["ID", "Title", "Type", "Priority", "Preconditions", "Test Data", "Steps", "Expected Result", "References"])

    for tc in test_cases:
        steps = tc.get("steps") or []
        preconditions = tc.get("preconditions") or []
        test_data = tc.get("test_data") or []
        references = tc.get("references") or []
        writer.writerow([
            tc.get("id", ""),
            tc.get("title", ""),
            tc.get("type", ""),
            tc.get("priority", ""),
            "\n".join(preconditions),
            "\n".join(test_data),
            "\n".join(steps),
            tc.get("expected_result", ""),
            "\n".join(references),
        ])

    timestamp = datetime.now().strftime("%Y-%m-%d-%H-%M")
    filename = f"testcases-{timestamp}.csv"
    # UTF-8 BOM so Google Sheets detects encoding correctly
    content = "﻿" + sio.getvalue()

    return Response(
        content=content.encode("utf-8"),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


app.add_route("/", chat_page, methods=["GET"])
app.add_route("/invocations", chat_page, methods=["GET"])
app.add_route("/chat", chat_api, methods=["POST"])
app.add_route("/ai/status", ai_status_api, methods=["GET"])
app.add_route("/knowledge", knowledge_list_api, methods=["GET"])
app.add_route("/knowledge", knowledge_create_api, methods=["POST"])
app.add_route("/knowledge/upload", knowledge_upload_api, methods=["POST"])
app.add_route("/knowledge/action", knowledge_action_api, methods=["POST"])
app.add_route("/export/csv", export_csv_api, methods=["POST"])
app.add_route("/config/types", config_types_api, methods=["GET"])


@app.entrypoint
def handler(payload: dict, context: RequestContext) -> dict:
    """Generate structured QC test cases from feature, requirement, and trained knowledge."""
    result = build_test_cases(payload)
    result["session_id"] = context.session_id
    return result


@app.ping
def health_check() -> PingStatus:
    """Health check for GET /health."""
    return PingStatus.HEALTHY


if __name__ == "__main__":
    app.run(port=8080, host="0.0.0.0")
